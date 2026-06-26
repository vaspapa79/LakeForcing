"""build_cover_letter.py -- regenerate paper/CoverLetter.docx as a resubmission
(revision-round) cover letter for the Computers & Geosciences major revision."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"C:/Users/vaspapa/Desktop/LakeForcing_OpenDrift")
OUT = ROOT / "paper" / "CoverLetter.docx"

doc = Document()
sec = doc.sections[0]
sec.left_margin = sec.right_margin = Inches(1.25)
sec.top_margin = Inches(1.0)
st = doc.styles['Normal']
st.font.name = 'Times New Roman'
st.font.size = Pt(11)


def para(text=None, align=None, after=6, before=0, bold_runs=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.space_before = Pt(before)
    if text is not None:
        if bold_runs:
            # text is a list of (segment, bold) tuples
            for seg, b in text:
                r = p.add_run(seg); r.bold = b; r.italic = False
        else:
            p.add_run(text)
    return p


J = WD_ALIGN_PARAGRAPH.JUSTIFY
R = WD_ALIGN_PARAGRAPH.RIGHT
L = WD_ALIGN_PARAGRAPH.LEFT

para('26 June 2026', align=R, after=18)
p = doc.add_paragraph(); p.alignment = L; p.paragraph_format.space_after = Pt(18)
p.add_run('To the Editors-in-Chief,'); p.add_run('\nComputers & Geosciences')
para('Dear Editors,', align=L, after=6)

para([('We are pleased to submit the revised version of our manuscript, ', False),
      ('"LakeForcing: a σ-to-z coupling algorithm and open pipeline for hydrodynamic and '
       'wind-wave forcing of inland lakes to drive Lagrangian transport models"', False),
      (', in response to the ', False), ('Major Revision', True),
      (' decision. We are grateful to the reviewer for an exceptionally careful and specific '
       'report; the recommendations were well taken, and addressing them has materially improved '
       'the transparency and reusability of the work. Every point — the five essential '
       'reproducibility items (R1–R5), the six major issues (M1–M6), the minor points '
       '(m1–m5), and the four questions (Q1–Q4) — has been addressed in the revised '
       'manuscript, in the repository, or in both, with no point left open.', False)],
     align=J, after=6, bold_runs=True)

para('A separate, point-by-point response (Response_to_Reviewers) accompanies this resubmission. '
     'In summary, the principal revisions are:', align=J, after=4)


def bullet(lead, rest):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    p.alignment = J
    r = p.add_run(lead); r.bold = True
    p.add_run(rest)
    return p


bullet('Full reproducibility of the toolchain. ',
       'We now specify exactly how the hydrodynamic engine is obtained (the official Deltares '
       'pre-built Delft3D 4 binary, release tag 4.07.01, Windows 11 64-bit — no local source '
       'build), the precise Copernicus CDS dataset and variable names used for the ERA5 retrieval '
       '(with credential and CDS-migration guidance), a pinned software environment '
       '(requirements.txt, conda-forge), and the provenance and assertions of the continuous-'
       'integration fixture that exercises the σ-to-z exporter from a clean checkout with no '
       'Delft3D install (R1, R2, R3, M6).')
bullet('Honest framing of the reference comparison. ',
       'Section 5.6 is reframed as an internal consistency check against an expert-built model from '
       'the same group — not an independent validation — and the abstract, overview and '
       'conclusions were aligned accordingly (M1).')
bullet('Strengthened quantitative reporting. ',
       'The conservation note now reports the median transport discrepancy across the depth range '
       '(it does not grow systematically in shallow basins), and the area–drift rank '
       'correlation is reported with and without the one hand-built lake (M2, m4).')
bullet('Algorithmic and methodological detail. ',
       'We state the SWAN output fields feeding the Stokes-drift equation, the Delft3D σ-layer '
       'variable and sign convention, the satellite-comparison methodology (binning, peak-hour '
       'selection, skin-to-bulk offset), and we expose the z-level set as a parameter (R4, M4, M5, '
       'm2).')

para([('We believe the work remains a strong fit for ', False),
      ('Computers & Geosciences', False),
      (' as original research at the interface of computational science and the geosciences: a '
       'reusable algorithm and an openly released, reproducible pipeline — centred on a fully '
       'specified σ-to-z coupling that bridges a terrain-following hydrodynamic engine to a '
       'generic Lagrangian tracker — demonstrated unchanged across twelve morphologically and '
       'climatically diverse lakes on all inhabited continents (36°S to 60°N). The '
       'generality is supported by a same-group model-to-model consistency check against an '
       'expert-built reservoir model and by an independent comparison of the exported surface '
       'temperature against satellite observations for four further lakes. The complete source code '
       'is released under the MIT licence with a continuous-integration test capsule and a pinned '
       'environment, and the generated twelve-lake forcing dataset is openly archived, so that the '
       'reported results are fully reproducible. We would again be glad to prepare a companion data '
       'descriptor for the forcing dataset as a linked co-submission (e.g. in ', False),
      ('Data in Brief', False),
      (') should the editors consider it appropriate.', False)],
     align=J, after=6, bold_runs=True)

para('The revised manuscript is original, has not been published previously, and is not under '
     'consideration for publication elsewhere. All authors have approved the revision and agree to '
     'its submission. The authors declare no competing interests. This research did not receive any '
     'specific grant from funding agencies in the public, commercial, or not-for-profit sectors; the '
     'work was carried out using the existing research infrastructure of CERTH-ITI. The source code '
     'is openly available at https://github.com/vaspapa79/LakeForcing and archived on Zenodo (concept '
     'DOI: https://doi.org/10.5281/zenodo.20627160); the generated twelve-lake forcing dataset is '
     'distributed as release assets of the same archive.', align=J, after=6)

para('Thank you for your continued consideration. We look forward to your response.', align=J, after=18)

para('Sincerely,', align=L, after=2)
para('Vassilios Papaioannou, on behalf of all co-authors', align=L, after=2)
para('Information Technologies Institute, Centre for Research and Technology Hellas (CERTH-ITI)',
     align=L, after=2)
para('6th km Charilaou-Thermi, 57001 Thessaloniki, Greece', align=L, after=2)
para('vaspapa@iti.gr  |  Tel. +30 697 285 4287', align=L, after=2)

doc.save(OUT)
print('wrote', OUT, '| paragraphs:', len(doc.paragraphs))
