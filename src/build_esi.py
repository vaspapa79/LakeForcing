"""
build_esi.py -- Earth Science Informatics (Springer) submission package -> paper2/.

Builds the manuscript docx from the SHARED source (paper/CAGEO_manuscript.md, so the
science stays in sync with the Computers & Geosciences package) plus an ESI-addressed
cover letter. No Highlights file (Springer does not use them).

  python src/build_esi.py
"""
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(r"C:/Users/vaspapa/Desktop/LakeForcing_OpenDrift")
sys.path.insert(0, str(ROOT / "src"))
import build_docx                      # noqa: E402
import build_aux_docs as aux           # noqa: E402  (reuse helpers + TITLE)

PAPER2 = ROOT / "paper2"
PAPER2.mkdir(exist_ok=True)
SERIF = "Times New Roman"
TITLE = aux.TITLE
L = WD_ALIGN_PARAGRAPH.LEFT
J = WD_ALIGN_PARAGRAPH.JUSTIFY


def build_manuscript():
    build_docx.MD = ROOT / "paper" / "CAGEO_manuscript.md"
    build_docx.DOCX = PAPER2 / "ESInformatics_manuscript.docx"
    build_docx.main()


def build_cover_letter():
    doc = Document()
    n = doc.styles["Normal"]; n.font.name = SERIF; n.font.size = Pt(11)
    n.paragraph_format.alignment = J

    def para(text, align=J, after=10):
        p = doc.add_paragraph(); p.alignment = align
        p.paragraph_format.space_after = Pt(after)
        for k, part in enumerate(aux.URL_RE.split(text)):   # URLs -> hyperlinks
            if not part:
                continue
            if k % 2 == 1:
                trail = ""
                while part and part[-1] in ".,;:":
                    trail = part[-1] + trail; part = part[:-1]
                aux.add_hyperlink(p, part, part)
                if trail:
                    p.add_run(trail)
            else:
                p.add_run(part)
        return p

    para("22 June 2026", align=L, after=6)
    para("To the Editors-in-Chief,\nEarth Science Informatics", align=L, after=12)
    para("Dear Editors,", align=L, after=10)

    para(f"We are pleased to submit our manuscript, “{TITLE}”, for "
         "consideration in Earth Science Informatics.")

    para("Lagrangian particle tracking is a standard tool for studying the transport of "
         "floating material, but in inland lakes its use is blocked by the absence of "
         "ready-made forcing: global ocean reanalyses stop at the coast, and lake "
         "hydrodynamic models are built by hand, one waterbody at a time. This is a "
         "growing gap, because recent cross-national surveys show that lakes and "
         "reservoirs are among the most acutely plastic-polluted freshwater systems on Earth.")

    para("Our manuscript presents LakeForcing, an open and reproducible Python pipeline "
         "that assembles bathymetry and meteorology from open global datasets, "
         "automatically configures and runs a coupled Delft3D-FLOW + Delft3D-WAVE (SWAN) "
         "simulation for an arbitrary lake, and exports CF-compliant NetCDF that drives "
         "Lagrangian particle trackers without modification. The methodological core is a "
         "fully specified sigma-to-z coupling algorithm that bridges Delft3D's "
         "terrain-following sigma-layers to fixed metric z-levels, together with the "
         "curvilinear-to-regular regridding, velocity rotation and surface Stokes-drift "
         "derivation needed to make lake hydrodynamics interoperable with a generic "
         "tracker. We demonstrate the pipeline, unchanged, across twelve morphologically "
         "and climatically diverse lakes on all inhabited continents, from 36°S to 60°N.")

    para("We believe the work fits the scope of Earth Science Informatics as original "
         "research at the interface of computer science and the geosciences: it "
         "contributes a reusable algorithm and an openly released, reproducible software "
         "pipeline, demonstrated generality across twelve lakes, and full software and "
         "data availability. It addresses a concrete and timely computational barrier to "
         "environmental transport modelling in freshwater systems. The generality is "
         "supported by a model-to-model benchmark against an expert-built reservoir model "
         "and by an independent comparison of the exported surface temperature against "
         "satellite observations for four further lakes. The complete source code is "
         "released under the MIT licence with a continuous-integration test capsule, and "
         "the generated twelve-lake forcing dataset is openly archived, so that the "
         "reported results are fully reproducible.")

    para("This manuscript is original, has not been published previously, and is not under "
         "consideration for publication elsewhere. All authors have approved the manuscript "
         "and agree to its submission. The authors declare no competing interests. This "
         "research did not receive any specific grant from funding agencies in the public, "
         "commercial, or not-for-profit sectors; the work was carried out using the existing "
         "research infrastructure of CERTH-ITI. The source code is openly available at "
         "https://github.com/vaspapa79/LakeForcing and archived on Zenodo (concept DOI: "
         "https://doi.org/10.5281/zenodo.20627160); the generated twelve-lake forcing "
         "dataset is distributed as release assets of the same archive.")

    para("Thank you for your consideration. We look forward to your response.", after=14)
    para("Sincerely,", align=L, after=2)
    para("Vassilios Papaioannou, on behalf of all co-authors", align=L, after=2)
    para("Information Technologies Institute, Centre for Research and Technology Hellas "
         "(CERTH-ITI)", align=L, after=2)
    para("6th km Charilaou-Thermi, 57001 Thessaloniki, Greece", align=L, after=2)
    para("vaspapa@iti.gr  |  Tel. +30 697 285 4287", align=L, after=2)

    out = PAPER2 / "CoverLetter.docx"; doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    build_manuscript()
    build_cover_letter()
