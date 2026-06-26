"""build_response_letter.py -- render paper/Response_to_Reviewers.docx (point-by-point
reply to the CAGEO strict review v01)."""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor

ROOT = Path(r"C:/Users/vaspapa/Desktop/LakeForcing_OpenDrift")
OUT = ROOT / "paper" / "Response_to_Reviewers.docx"

doc = Document()
st = doc.styles['Normal']; st.font.name = 'Calibri'; st.font.size = Pt(10.5)
NAVY = RGBColor(0x1f, 0x3b, 0x63)


def H(t, sz=14, before=10, after=6):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True
    r.font.size = Pt(sz); r.font.color.rgb = NAVY
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def P(t, italic=False, before=0, after=6):
    p = doc.add_paragraph(); r = p.add_run(t); r.italic = italic
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    return p


def item(tag, concern, response):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(tag); r.bold = True; r.font.color.rgb = NAVY
    q = doc.add_paragraph(); q.paragraph_format.space_after = Pt(2)
    rr = q.add_run('Reviewer: '); rr.bold = True; rr.italic = True
    q.add_run(concern).italic = True
    a = doc.add_paragraph(); a.paragraph_format.space_after = Pt(6)
    ra = a.add_run('Response: '); ra.bold = True
    a.add_run(response)


H('Response to Reviewer Comments', 16, before=0, after=4)
P('Manuscript: "LakeForcing - a sigma-to-z coupling algorithm and open pipeline for hydrodynamic '
  'and wind-wave forcing of inland lakes to drive Lagrangian transport models"', italic=True)
P('Journal: Computers & Geosciences    |    Decision: Major Revision', italic=True)
P('We thank the reviewer for a careful, constructive and unusually specific review. The concerns '
  'were well taken and every one has been addressed in the revised manuscript and the repository. '
  'Below we reply point by point; reviewer text is paraphrased in italics, and manuscript locations '
  'refer to the revised version. Key facts added to the paper (Delft3D acquisition, CDS '
  'dataset/variable strings, SWAN field names, fixture provenance, conservation statistics across the '
  'depth range, and the with/without-Polyfytos correlations) were taken directly from the released '
  'code and data.', before=4)

H('1. Essential - reproducibility blockers')
item('R1 - Delft3D version and build provenance',
     'Delft3D 4 is a legacy branch with a non-trivial build; no container/binary or build '
     'configuration is given, so the hydrodynamic runs cannot be reproduced.',
     'We now state in Section 3.2 and the Software-availability table that we did not build Delft3D '
     'from source: the reported runs used the official Deltares pre-built Delft3D 4 binary (release '
     'tag 4.07.01) from the Deltares open-source portal (oss.deltares.nl/web/delft3d), executed on '
     'Windows 11 (64-bit) through the distribution\'s own launch scripts. Using the vendor binary '
     'means the FLOW/WAVE results do not depend on a user\'s compiler, MPI variant or patch level; '
     'the exact tag, OS and download reference are recorded in the README and the table. We also '
     'emphasise that the most reusable component - the sigma-to-z exporter - is made independently '
     'reproducible without the engine via the CI fixture (R3).')
item('R2 - ERA5/CDS retrieval specification',
     'No CDS dataset names, variable short-names, product types, cdsapi version, or credential setup '
     'are given; the CDS migration will break the retrieval.',
     'Added to Section 3.2 and the README: all three get_era5_*.py modules call cdsapi against '
     'dataset "reanalysis-era5-single-levels" with product_type "reanalysis" and the exact '
     'short-names - 10m_u_component_of_wind, 10m_v_component_of_wind (wind); 2m_temperature, '
     '2m_dewpoint_temperature, total_cloud_cover, surface_solar_radiation_downwards (heat flux); '
     'total_precipitation, evaporation, 2m_temperature (mass balance). We note the legacy-CDS '
     'retirement, give the confirmed cdsapi version and the ~/.cdsapirc credential steps in the '
     'README, and state that the variable names are unchanged across the migration (only the '
     'endpoint/key differ).')
item('R3 - test fixture provenance',
     'The CI fixture is undescribed (which lake, grid size, duration, how generated, what the test '
     'checks).',
     'The fixture is now fully described in Section 3.2 and the Software-availability table: it is a '
     'subset of the smallest demonstration lake, Erken (EPSG:32634); trim-erken_mini.nc keeps the '
     'first two time steps and the variables the exporter reads on the 125x52 grid with 14 '
     'sigma-layers, with the matching wavm-erken_mini.nc. Both are produced from the full Erken run '
     'by the versioned tests/build_fixture.py and are committed and checksummed. We list what '
     'tests/test_cf_export.py asserts quantitatively (clean exit, CF Conventions, presence and CF '
     'standard_names of u/v/temp/zeta/Hs, monotonic z-levels with surface at 0 and sub-bed masking, '
     'monotonic lon/lat, finite wet surface).')
item('R4 - SWAN output variables in Eq. (6)',
     'The SWAN fields used for Hs, Tp and direction are not stated; choosing the wrong period gives a '
     'factor-of-two Stokes error.',
     'Specified in Section 2.5(4): Hs is taken from HSIGN, Tp from RTP (SWAN\'s relative peak period '
     'at the spectral peak - explicitly not an integral mean such as TM01/TMM10), and direction from '
     'DIR, all from the wavm-*.nc map file and requested in the .mdw output block written by the wave '
     'generator. We add the explicit warning that substituting a mean period for RTP would bias the '
     'Stokes magnitude by a factor of order two.')
item('R5 - Zenodo archive contents',
     'The manuscript must let the reviewer verify that the deposit actually contains the twelve-lake '
     'NetCDF dataset, not only code; if only code is deposited, correct the Data-availability '
     'statement and deposit the data separately.',
     'On checking the deposit we confirmed the reviewer\'s suspected case: the Zenodo record is '
     'created through the GitHub-release integration and archives the source-code snapshot only (the '
     'generated NetCDF files are large and licensed CC-BY-4.0, and are not committed to the repo). We '
     'have therefore corrected the Data-availability statement rather than overstate the Zenodo '
     'contents: it now says the source code is archived on Zenodo (concept DOI, MIT), while the '
     'twelve-lake forcing dataset (CC-BY-4.0; twelve <lake>_forcing.nc, ~0.9 GB uncompressed / ~0.7 GB '
     'as LakeForcing-OpenDrift_forcing_dataset_v1.0.0.zip) is distributed as a versioned GitHub '
     'release asset with a manifest (DATASET_MANIFEST.txt) listing each file\'s size and SHA-256 '
     'checksum for verification against Section 4. We are glad to additionally mint a separate Zenodo '
     'data DOI for the forcing dataset if the editors prefer a single citable data archive.')

H('2. Major issues')
item('M1 - Polyfytos benchmark circularity',
     'The reference (Papaioannou et al., 2025) is the authors\' own prior model; presenting it as '
     'independent validation overstates its weight.',
     'Agreed. Section 5.6 is retitled "Consistency check against an expert-built reference model from '
     'the same group" and its opening now states plainly that the reference is the authors\' own '
     'previously published model and that the comparison is an internal consistency check, not '
     'independent validation. The abstract, the Section-5 overview and the Conclusions were reworded '
     'the same way (the abstract now says "in a consistency check against an expert-built reference '
     'model ... it reproduces that model\'s surface temperature"). The directional-correlation '
     'failure (|rho| ~ 0.10, unchanged at ~0.09 with river discharge added) is retained and discussed '
     'as delimiting what the automated configuration does and does not reproduce.')
item('M2 - Conservation note, single lake only',
     'The 2% transport discrepancy is given for one unnamed lake; shallow lakes should be worse; '
     'report range / worst case and clarify the statistic.',
     'The conservation note in Section 2.5 now reports the median relative depth-integrated transport '
     'error across the depth range and names the representative lake (Bornos, ~2.1%). New numbers: '
     '~2.0% for the shallow Sea of Galilee (~6 m), ~0.5% for shallow Poyang (~6 m), ~0.9% for deep '
     'Mead (~40 m). The key finding is that the median does NOT grow systematically as the basin '
     'shoals. We also clarify that the per-column mean is larger (up to ~10% in some lakes) but is '
     'dominated by the slowest cells, where a relative error measure is ill-conditioned (vanishing '
     'denominator) while the absolute error stays small.')
item('M3 - Table 3 does not flag Polyfytos',
     'Polyfytos (hand-built) sits in the same style as the eleven auto-generated lakes and may '
     'mislead.',
     'Table 3 now renders the Polyfytos row in bold and labels the lake cell "Polyfytos* '
     '(hand-built)"; the caption and footnote state that all twelve lakes are auto-generated except '
     'Polyfytos, which is the export-path control and not an auto-generated case.')
item('M4 - sigma-layer indexing convention',
     'It is not stated whether sigma_k is read directly from Delft3D or transformed; a sign/indexing '
     'error would silently invert the profile.',
     'Section 2.5(1) now specifies that sigma_k is read directly from the SIG_LYR variable of the '
     'trim-*.nc map file (with S1, DPS and ALFAS), that no vertical offset is applied, and that the '
     'only transformation is a single sign normalisation for builds that store the fractions '
     'positive-down on [0,1], so Eq. (4) always receives sigma=0 at the surface and sigma=-1 at the '
     'bed, with the surface-most layer mapped to the top z-level.')
item('M5 - Satellite comparison methodology',
     'Missing: binning resolution, aggregation method, skin-to-bulk correction, and how the peak hour '
     'was chosen.',
     'Section 5.7 now states: the Landsat thermal band is loaded in geographic coordinates at '
     '~30-60 m (basin-dependent) and aggregated onto the model grid by area-mean binning (not '
     'nearest-neighbour); the comparison hour is the model\'s own diurnal basin-mean maximum (close '
     'to the early-afternoon overpass), with the satellite UTC stamp shifted into the lake\'s local '
     'zone before matching; and no skin-to-bulk correction was applied - we note the satellite skin '
     'runs of order 0.2-0.5 K from the model bulk, small against the 4-5 C warm-lake offset but '
     'comparable to the sub-0.2 C winter-control agreement and therefore not to be over-interpreted '
     'there.')
item('M6 - Pinned environment file',
     'No environment.yml / pinned requirements; minor-version drift can break results.',
     'A pinned requirements.txt is present at the repository root and is now referenced in Section 3.2 '
     'and the Software-availability table. It pins opendrift==1.14.9 and recommends a conda-forge '
     'Python 3.11 environment so that rasterio/pyproj bind to a consistent GDAL/PROJ build (the usual '
     'cross-platform breakage point); the remaining pure-Python packages tolerate patch-level drift. '
     'Python 3.11 and the tested OS (Windows 11 64-bit; also Linux 64-bit) are stated.')

H('3. Minor issues')
item('m1 - Eq. (8) cos(phi) placement',
     'The cos(phi) factor appears inside the squared longitude term and may be dimensionally '
     'inconsistent.',
     'We checked the typeset equation against the code (make_equations.py / omml_equations.py and the '
     'demo distance metric). The form is correct: D = sqrt( [(lon_i-lon_0)*cos(phi)*Re]^2 + '
     '[(lat_i-lat_0)*Re]^2 ). The cos(phi) multiplies only the longitude difference, inside its own '
     'bracket, before scaling by Re (km/deg) - this is the standard equirectangular approximation and '
     'is dimensionally consistent (deg x dimensionless x km/deg = km). No change was required; we have '
     're-verified the rendering.')
item('m2 - z-level set hard-coded',
     'The output z-levels are not exposed as a parameter; a deep/benthic application cannot change '
     'them without editing source.',
     'Section 2.5 now states that the target set is a single exposed parameter (Z_LEVELS in '
     'cf_export.py, overridable from the exporter\'s command line), changeable in one place for deep '
     'or near-bed applications (e.g. Lake Baikal) without altering the reconstruction logic.')
item('m3 - Humidity accuracy overstated',
     'The 0.4% figure is the saturation-vapour-pressure error, not the RH error.',
     'Section 2.3 now states the 0.4% is the saturation-vapour-pressure error optimised by Alduchov '
     'and Eskridge (1996), and adds that the RH error - a ratio of two such approximations - can '
     'exceed this under large dew-point depressions, while remaining small for the near-surface lake '
     'conditions sampled here.')
item('m4 - Spearman robustness to Polyfytos',
     'With n=12, report area-vs-drift rho and p with and without the Polyfytos outlier.',
     'Section 5.3 and the Conclusions now report that excluding Polyfytos the primary area-drift '
     'correlation weakens from rho=+0.63 (p=0.03) to rho=+0.54 (p=0.09) - same sign and magnitude, '
     'but no longer significant at 5% - so the twelve-lake significance is partly carried by the '
     'hand-built outlier; the relation is presented as a coherent tendency rather than a law. The '
     'other predictors are essentially unchanged by the exclusion (reported in Section 5.3).')
item('m5 - Parcels compatibility claim',
     'Either remove the Parcels claim or add a tested Parcels run and document the version.',
     'We chose to confine the claim to a stated expectation / future-work note. Section 5.5 now says '
     'only the OpenDrift path (1.14.9) is demonstrated and tested; Parcels compatibility is expected '
     'by construction from the CF interface but is explicitly unverified, and Parcels is not listed as '
     'a tested dependency in the Software-availability table.')

H('4. Questions for the authors')
item('Q1 - SWAN grid for Polyfytos',
     'For the clockwise hand-built FLOW grid, does SWAN use the pipeline\'s CCW grid, and does mapping '
     'back introduce extra interpolation error?',
     'SWAN runs on the pipeline\'s own counter-clockwise grid. There is no back-mapping onto the FLOW '
     'domain: the exporter regrids the FLOW fields and the SWAN fields independently onto the common '
     'regular lon/lat raster of Section 2.5, so each field undergoes only the single '
     'curvilinear/CCW-to-regular step already bounded there, and no additional interpolation error is '
     'incurred. We added this clarification to Section 5.6.')
item('Q2 - 35 C cap and transport statistics',
     'Does the temperature cap affect any lake\'s drift; are reported drifts post-cap?',
     'The reported drifts are computed on the exported (post-cap) fields, but they are unaffected by '
     'the cap: advection is driven by currents, surface Stokes drift and wind - none of which the cap '
     'touches - not by the temperature field. The cap only suppresses isolated thin-cell shoreline '
     'temperature spikes. Added explicitly to Section 5.4.')
item('Q3 - Manuscript regenerated by the pipeline',
     'Does build_docx.py generate the submitted .docx; can the equations/figures be rebuilt '
     'bit-identically?',
     'Yes. The submitted .docx is the output of build_docx.py rendering paper/CAGEO_manuscript.md; '
     'equations are emitted as native Office Math (OMML) from omml_equations.py and the figures are '
     'the files written by make_figures.py, so the displayed equations and figures are by '
     'construction identical to the pipeline\'s output. The README documents the single rebuild '
     'command. Noted in Section 3.1.')
item('Q4 - River-discharge feature documented/tested',
     'Is the optional river-discharge boundary documented and tested in the repository?',
     'River discharge is supported as an optional open-boundary condition (used for the Polyfytos '
     'comparison, configured via _build_river_config.py). It is exercised in that benchmark but is not '
     'covered by the CI test; this is now flagged as an "Optional features" row in the '
     'Software-availability table and stated in Section 3.2.')

H('5. Editorial notes')
P('All editorial suggestions were adopted: (i) the abstract now attributes the 0.85 C RMSE to the '
  'expert-built reference model rather than to observations; (ii) Table 1\'s tracker-reader row was '
  'reworded to "Tracker native readers (e.g., ROMS-format, generic CF) - no Delft3D reader available; '
  'sigma-fields not directly ingestible", avoiding any implication that OpenDrift is deficient; '
  '(iii) Section 2.4 and Table 2 now read "Dalton number 0.0013 and Stanton number 0.0013" (both '
  'dimensionless bulk transfer coefficients); (iv) Section 5.2 "make ... and exercise" -> "makes ... '
  'and exercises". On equation typesetting: in the source pipeline the display equations are numbered '
  'Office Math (OMML) objects rendered by omml_equations.py, not table cells, so they are already '
  'proper, sequentially numbered equation objects; this transfers directly to the Elsevier template.')

P('We believe these revisions fully resolve the Essential and Major points and close every item on '
  'the reproducibility checklist. We are grateful for the review, which has materially improved the '
  'transparency and reusability of the work.', before=8)

doc.save(OUT)
print('wrote', OUT, '| paragraphs:', len(doc.paragraphs))
